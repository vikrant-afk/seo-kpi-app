"""Generate a Week-on-Week (WoW) SEO RCA as a Word (.docx) document.

The KPI *numbers* come straight from the KPI rows the app already built (accurate,
never invented). The *narrative* — analysis, root-cause reasoning, actions and
way-forward — is written by Claude via the Anthropic API, reusing the same
anthropic key/model the app already has in settings.
"""
from __future__ import annotations

import io
import re

import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ANTHROPIC_URL = "https://api.anthropic.com/v1/messages"
ACCENT = RGBColor(0x0E, 0x6E, 0x4E)   # green (matches the app)
INK = RGBColor(0x14, 0x1A, 0x16)
MUTED = RGBColor(0x5B, 0x61, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# --------------------------------------------------------------------------- Claude
def _call_claude(prompt: str, cfg: dict, max_tokens: int = 4096) -> str:
    key = cfg.get("anthropic_key")
    model = cfg.get("anthropic_model") or "claude-3-5-sonnet-latest"
    if not key:
        raise RuntimeError("Anthropic key not configured (cfg['anthropic_key']).")
    r = requests.post(
        ANTHROPIC_URL, timeout=180,
        headers={"x-api-key": key, "anthropic-version": "2023-06-01",
                 "content-type": "application/json"},
        json={"model": model, "max_tokens": max_tokens,
              "messages": [{"role": "user", "content": prompt}]},
    )
    r.raise_for_status()
    data = r.json()
    return "".join(b.get("text", "") for b in data.get("content", [])
                   if b.get("type") == "text").strip()


def _kpi_lines(rows) -> str:
    out = []
    for r in rows:
        line = f"- {r['metric']}: current={r['current']}"
        if r.get("previous") not in (None, "—"):
            line += f", previous={r['previous']}"
        if r.get("change_pct") is not None:
            line += f", change={r['change_pct']}%"
        line += f" (source: {r['source']})"
        out.append(line)
    return "\n".join(out)


def _prompt(rows, meta, comparison_on) -> str:
    window = meta.get("range", "")
    has_cmp = "present" if comparison_on else "NOT present (single period only)"
    return f"""You are a senior SEO analyst writing a Week-on-Week (WoW) Search Performance / RCA report for the website {meta.get('domain')}.
Reporting window: {window}. Week-over-week comparison data is {has_cmp}.

Use ONLY the KPI figures below. Do not invent any numbers. When you cite a movement, use the exact % change shown.
{_kpi_lines(rows)}

Write the report BODY as clean Markdown using these exact section headings (level-2, i.e. "## "), in this order:

## Executive Summary
3-5 sentences: the headline WoW story and overall direction.

## KPI Analysis
Interpret traffic (total users, organic sessions, LLM traffic) and what the movements imply.

## Query Click & Impression Growth
Analyse overall vs non-branded (NB) clicks, impressions and CTR. Explain what the branded/non-branded split suggests.

## Authority
Comment on Domain Rating (DR) and what it signals for the link profile this period.

## Technical
Comment on desktop/mobile PageSpeed and any technical risk to rankings.

## Actions Taken
3-6 bullets of plausible SEO work done this cycle that would drive these results.

## Way Forward
4-6 prioritised bullet recommendations for next week.

Rules: Be specific and quantify with the numbers above. Where deeper data (keyword-level rankings, query/page breakdowns) would be required to be certain of a root cause, say so briefly instead of inventing figures. Do NOT restate the KPI table and do NOT add a title — start directly at "## Executive Summary"."""


# --------------------------------------------------------------------------- docx helpers
def _shade(cell, hex_fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = ACCENT if level == 1 else INK
    run.font.size = Pt(14 if level == 1 else 12)
    return p


def _add_runs(paragraph, text):
    text = text.replace("`", "")
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _add_markdown(doc, md):
    for raw in md.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            _heading(doc, line[4:].strip(), level=2)
        elif line.startswith("## "):
            _heading(doc, line[3:].strip(), level=1)
        elif line.startswith("# "):
            _heading(doc, line[2:].strip(), level=1)
        elif re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
        elif re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\s*\d+\.\s+", "", line))
        else:
            _add_runs(doc.add_paragraph(), line)


def _fmt_change(cp):
    if cp is None:
        return "—"
    if cp == float("inf"):
        return "NEW"
    arrow = "▲" if cp > 0 else ("▼" if cp < 0 else "•")
    return f"{arrow} {abs(cp)}%"


def _kpi_table(doc, rows, comparison_on):
    headers = (["Metric", "Current", "Previous", "Change", "Source"] if comparison_on
               else ["Metric", "Current", "Source"])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        _shade(cell, "0E6E4E")
    for r in rows:
        vals = ([r["metric"], str(r["current"]), str(r.get("previous", "—")),
                 _fmt_change(r.get("change_pct")), r["source"]] if comparison_on
                else [r["metric"], str(r["current"]), r["source"]])
        cells = table.add_row().cells
        for i, v in enumerate(vals):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(v)
            run.font.size = Pt(9)
    return table


# --------------------------------------------------------------------------- public
def generate_rca_docx(rows, meta, cfg, comparison_on) -> bytes:
    """Return a .docx (as bytes) of the WoW RCA. Raises on Claude/API failure."""
    body_md = _call_claude(_prompt(rows, meta, comparison_on), cfg)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # Title block
    t = doc.add_paragraph()
    tr = t.add_run("Search Performance Growth Report — WoW RCA")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = INK

    sub = doc.add_paragraph()
    sr = sub.add_run(f"{meta.get('domain', '')}    ·    {meta.get('range', '')}")
    sr.font.size = Pt(10.5)
    sr.font.color.rgb = MUTED

    gen = doc.add_paragraph()
    gr = gen.add_run(f"Generated {meta.get('generated_at', '')}")
    gr.italic = True
    gr.font.size = Pt(9)
    gr.font.color.rgb = MUTED

    doc.add_paragraph()
    _heading(doc, "KPI Matrix", level=1)
    _kpi_table(doc, rows, comparison_on)
    doc.add_paragraph()

    _add_markdown(doc, body_md)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run("Domain Rating by Ahrefs (https://ahrefs.com/). "
                      "Sources: GA4, Google Search Console, PageSpeed Insights, Ahrefs.")
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- Claude Desktop handoff
def build_claude_desktop_brief(rows, meta, comparison_on) -> str:
    """Return a ready-to-paste brief for Claude Desktop.

    The Streamlit app cannot reach your Claude Desktop connectors or skills
    (those live in the desktop client, not on this server). So instead of trying
    to trigger them, this builds a brief you paste into Claude Desktop, where the
    Google Analytics / Search Console connectors and the wow-rca-report skill
    actually run.
    """
    window = meta.get("range", "")
    cmp_note = ("A week-over-week comparison window IS included (current vs previous)."
                if comparison_on else
                "Only a single period is included — no comparison window was selected.")
    return f"""# Week-on-Week SEO RCA — brief for Claude Desktop

**Client / site:** {meta.get('domain', '')}
**Reporting window:** {window}
{cmp_note}

## What I need you to do
1. Using my connected **Google Analytics (GA4)** and **Google Search Console** connectors in Claude Desktop (use these directly — do NOT use Supermetrics or any other aggregator), pull the same metrics listed below for this reporting window (and the comparison window if shown).
2. **Cross-verify** those live figures against the tool-extracted numbers below. For any metric that differs by more than ~2%, call it out and give the likely reason (data sampling, timezone, branded/non-branded filter, or a GA4/GSC property mismatch).
3. Then produce the full **Week-on-Week Search Performance Growth / RCA report** using the **wow-rca-report skill**, using the verified numbers. Where a verified value differs from the tool value, prefer the connector-pulled value and note the correction.

## Tool-extracted KPI figures (verify against these)
{_kpi_lines(rows)}

## Notes
- "NB" = non-branded (a regex excludes branded queries from GSC).
- DR (Domain Rating) is from Ahrefs' free public endpoint. DA, Spam Score and Referring Domains were removed from this tool.
- Use only my Google Analytics and Google Search Console connectors for verification. If a metric can't be matched through those connectors, say so explicitly rather than estimating.
"""
